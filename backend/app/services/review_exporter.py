"""Generate a Word document review of a contract's analysis.

Produces a self-contained DOCX summarising:
  - Overall risk level and AI-generated summary
  - Each yellow/red clause with original text, plain-English
    explanation, market comparison, and suggested revision
  - Green clauses listed compactly at the end

Typography uses Aptos (Microsoft's modern default sans-serif, shipping
with Word 2024+). Older Word installs substitute Calibri, which still
renders the document cleanly.

The exporter takes pre-fetched ORM objects rather than a DB session so
it stays trivially testable and doesn't couple document generation to
async I/O.
"""

# pyright: reportAttributeAccessIssue=false

from __future__ import annotations

import logging
import re
from collections.abc import Iterable
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.models.clause import Clause
from app.models.contract import Contract

logger = logging.getLogger(__name__)

# Typography
FONT_BODY = "Aptos"
FONT_DISPLAY = "Aptos Display"
FONT_MONO = "Aptos Mono"

# Spacing constants. Word measures spacing in points (Pt), not pixels.
# These values were tuned to give the document an editorial feel:
# generous breathing room around section breaks, tighter inside each
# clause so related content feels grouped.
SPACE_BEFORE_H1 = Pt(24)
SPACE_AFTER_H1 = Pt(12)
SPACE_BEFORE_H2 = Pt(18)
SPACE_AFTER_H2 = Pt(6)
SPACE_BEFORE_H3 = Pt(10)
SPACE_AFTER_H3 = Pt(4)
SPACE_AFTER_BODY = Pt(6)

# Risk indicators. Picked for legibility on white in Word's default theme.
RISK_COLOR_HIGH = RGBColor(0xB4, 0x1C, 0x1C)
RISK_COLOR_MED = RGBColor(0xB4, 0x7A, 0x1C)
RISK_COLOR_LOW = RGBColor(0x2C, 0x6E, 0x3A)
MUTED_COLOR = RGBColor(0x6B, 0x6B, 0x6B)
RULE_COLOR = RGBColor(0xCC, 0xCC, 0xCC)

RISK_LABELS = {
    "red": "High risk",
    "yellow": "Watch",
    "green": "No concerns flagged",
}

OVERALL_RISK_LABELS = {
    "high": "High",
    "medium": "Medium",
    "low": "Low",
}

OVERALL_RISK_COLORS = {
    "high": RISK_COLOR_HIGH,
    "medium": RISK_COLOR_MED,
    "low": RISK_COLOR_LOW,
}


def build_review_docx(contract: Contract, clauses: Iterable[Clause]) -> BytesIO:
    """Build a ClauseGuard review DOCX in memory.

    Args:
        contract: The contract being reviewed.
        clauses: All clauses for the contract. Will be sorted by
            position and deduplicated by trimmed original text.

    Returns:
        BytesIO positioned at byte 0, ready to stream.
    """
    deduped = _dedupe_clauses(clauses)
    sorted_clauses = sorted(deduped, key=lambda c: c.position)

    concerns = [c for c in sorted_clauses if c.risk_level in ("yellow", "red")]
    benign = [c for c in sorted_clauses if c.risk_level == "green"]

    doc = Document()
    _patch_default_template(doc)
    _configure_page(doc)
    _configure_styles(doc)
    _configure_header_footer(doc, contract)

    _write_cover(doc, contract)
    _write_overall_assessment(
        doc, contract, concerns_count=len(concerns), benign_count=len(benign)
    )
    _write_concerns(doc, concerns)
    _write_benign(doc, benign)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _patch_default_template(doc: DocumentType) -> None:
    """python-docx's default template ships a malformed `<w:zoom>`
    element in settings.xml: the OOXML schema requires `w:percent`,
    but the template only sets `w:val="bestFit"`. Word recovers but
    warns 'Word found unreadable content' on open. Add the missing
    attribute so the document validates cleanly.

    This patches the bundled python-docx template, not anything we
    write. Can be removed if python-docx ships a fix upstream.
    """
    settings = doc.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is not None and qn("w:percent") not in zoom.attrib:
        zoom.set(qn("w:percent"), "100")


# ---------------------------------------------------------------------------
# Setup: page geometry, styles, headers and footers
# ---------------------------------------------------------------------------


def _configure_page(doc: DocumentType) -> None:
    """US Letter with generous margins. The slightly-wider-than-default
    margins give body paragraphs a comfortable measure (~75ch at 11pt
    Aptos) without resorting to a separate text frame."""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)


def _configure_styles(doc: DocumentType) -> None:
    """Tune the built-in styles. python-docx exposes Normal, Heading 1
    through 9, Title, Subtitle, List Bullet, and others; we configure
    the ones we actually use."""
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    _set_complex_script_font(normal, FONT_BODY)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    normal.paragraph_format.space_after = SPACE_AFTER_BODY
    normal.paragraph_format.line_spacing = 1.4

    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT_DISPLAY
    _set_complex_script_font(h1, FONT_DISPLAY)
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    h1.paragraph_format.space_before = SPACE_BEFORE_H1
    h1.paragraph_format.space_after = SPACE_AFTER_H1
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT_DISPLAY
    _set_complex_script_font(h2, FONT_DISPLAY)
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    h2.paragraph_format.space_before = SPACE_BEFORE_H2
    h2.paragraph_format.space_after = SPACE_AFTER_H2
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    h3.font.name = FONT_BODY
    _set_complex_script_font(h3, FONT_BODY)
    h3.font.size = Pt(10)
    h3.font.bold = True
    h3.font.all_caps = True
    h3.font.color.rgb = MUTED_COLOR
    h3.paragraph_format.space_before = SPACE_BEFORE_H3
    h3.paragraph_format.space_after = SPACE_AFTER_H3
    h3.paragraph_format.keep_with_next = True

    title = doc.styles["Title"]
    title.font.name = FONT_DISPLAY
    _set_complex_script_font(title, FONT_DISPLAY)
    title.font.size = Pt(32)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = FONT_BODY
    _set_complex_script_font(subtitle, FONT_BODY)
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = MUTED_COLOR
    subtitle.font.italic = False
    subtitle.paragraph_format.space_after = Pt(2)

    bullet = doc.styles["List Bullet"]
    bullet.font.name = FONT_BODY
    _set_complex_script_font(bullet, FONT_BODY)
    bullet.font.size = Pt(11)
    bullet.paragraph_format.space_after = Pt(2)
    bullet.paragraph_format.line_spacing = 1.3


def _set_complex_script_font(style, font_name: str) -> None:
    """python-docx sets the Latin font when you assign `font.name`,
    but Word also tracks an east-Asian and complex-script font that
    can fall back to Calibri if not also set. This nudges them all to
    the same family so Aptos renders consistently."""
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def _configure_header_footer(doc: DocumentType, contract: Contract) -> None:
    """Page header carries the contract name as a quiet caption.
    Footer carries the page number on the right and 'ClauseGuard
    review' on the left, both in muted small caps to feel ambient
    rather than insistent."""
    section = doc.sections[0]

    header_para = section.header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header_para.add_run(contract.file_name)
    header_run.font.name = FONT_BODY
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = MUTED_COLOR

    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    left_run = footer_para.add_run("ClauseGuard review")
    left_run.font.name = FONT_BODY
    left_run.font.size = Pt(9)
    left_run.font.color.rgb = MUTED_COLOR

    # Tab to push the page number to the right margin.
    tab_run = footer_para.add_run("\t")
    tab_run.font.size = Pt(9)

    page_field_run = footer_para.add_run()
    page_field_run.font.name = FONT_BODY
    page_field_run.font.size = Pt(9)
    page_field_run.font.color.rgb = MUTED_COLOR
    _add_page_number_field(page_field_run)


def _add_page_number_field(run) -> None:
    """python-docx doesn't expose a high-level page-number element, so
    we inject the OOXML field directly. Three children: fldChar begin,
    instrText with 'PAGE', fldChar end."""
    fld_begin = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    instr = run._element.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
    instr.text = "PAGE"
    fld_end = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_end)


# ---------------------------------------------------------------------------
# Content sections
# ---------------------------------------------------------------------------


def _write_cover(doc: DocumentType, contract: Contract) -> None:
    """Cover block: large title, contract filename in subtitle weight,
    analysed date in muted small text. No section break — the rest of
    the document flows from here."""
    eyebrow = doc.add_paragraph()
    eyebrow.paragraph_format.space_after = Pt(2)
    eyebrow_run = eyebrow.add_run("Contract review")
    eyebrow_run.font.name = FONT_BODY
    eyebrow_run.font.size = Pt(10)
    eyebrow_run.font.bold = True
    eyebrow_run.font.all_caps = True
    eyebrow_run.font.color.rgb = MUTED_COLOR

    doc.add_paragraph("ClauseGuard Review", style="Title")
    doc.add_paragraph(contract.file_name, style="Subtitle")

    analyzed_at = contract.analyzed_at or contract.created_at
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_after = Pt(12)
    date_run = date_para.add_run(f"Analysed {_format_date(analyzed_at)}")
    date_run.font.name = FONT_BODY
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = MUTED_COLOR


def _write_overall_assessment(
    doc: DocumentType,
    contract: Contract,
    *,
    concerns_count: int,
    benign_count: int,
) -> None:
    doc.add_heading("Overall assessment", level=1)

    # Risk + counts on a single line, separated visually.
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)

    risk_label = OVERALL_RISK_LABELS.get(contract.overall_risk or "", "Not assessed")
    label_run = meta.add_run("Risk level   ")
    label_run.font.name = FONT_BODY
    label_run.font.size = Pt(10)
    label_run.font.color.rgb = MUTED_COLOR
    label_run.font.all_caps = True
    label_run.bold = True

    risk_run = meta.add_run(risk_label)
    risk_run.font.name = FONT_BODY
    risk_run.font.size = Pt(11)
    risk_run.bold = True
    if contract.overall_risk in OVERALL_RISK_COLORS:
        risk_run.font.color.rgb = OVERALL_RISK_COLORS[contract.overall_risk]

    if contract.contract_type:
        sep = meta.add_run("     ")
        sep.font.size = Pt(10)

        type_label_run = meta.add_run("Type   ")
        type_label_run.font.name = FONT_BODY
        type_label_run.font.size = Pt(10)
        type_label_run.font.color.rgb = MUTED_COLOR
        type_label_run.font.all_caps = True
        type_label_run.bold = True

        type_value_run = meta.add_run(contract.contract_type.upper())
        type_value_run.font.name = FONT_BODY
        type_value_run.font.size = Pt(11)

    # Counts caption.
    counts_para = doc.add_paragraph()
    counts_para.paragraph_format.space_after = Pt(14)
    counts_text = (
        f"{concerns_count} clause{'s' if concerns_count != 1 else ''} flagged for attention"
    )
    if benign_count:
        counts_text += f", {benign_count} reviewed without issue."
    else:
        counts_text += "."
    counts_run = counts_para.add_run(counts_text)
    counts_run.font.name = FONT_BODY
    counts_run.font.size = Pt(10)
    counts_run.font.color.rgb = MUTED_COLOR

    if contract.summary:
        summary_para = doc.add_paragraph(_normalize_text(contract.summary))
        summary_para.paragraph_format.space_after = Pt(8)


def _write_concerns(doc: DocumentType, concerns: list[Clause]) -> None:
    doc.add_heading("Concerns", level=1)

    if not concerns:
        doc.add_paragraph(
            "No concerns flagged. The contract was reviewed in full and every clause "
            "tracks market-standard language."
        )
        return

    for index, clause in enumerate(concerns):
        _write_clause(doc, clause, is_last=index == len(concerns) - 1)


def _write_clause(doc: DocumentType, clause: Clause, *, is_last: bool) -> None:
    type_label = _format_clause_type(clause.clause_type)
    position_label = str(clause.position).zfill(3)

    doc.add_heading(f"Clause {position_label}   {type_label}", level=2)

    # Risk caption sits immediately under the heading, small, coloured.
    risk_para = doc.add_paragraph()
    risk_para.paragraph_format.space_after = Pt(10)
    risk_run = risk_para.add_run(RISK_LABELS.get(clause.risk_level or "", "Pending analysis"))
    risk_run.font.name = FONT_BODY
    risk_run.font.size = Pt(10)
    risk_run.bold = True
    risk_run.font.all_caps = True
    if clause.risk_level == "red":
        risk_run.font.color.rgb = RISK_COLOR_HIGH
    elif clause.risk_level == "yellow":
        risk_run.font.color.rgb = RISK_COLOR_MED

    _write_subsection(doc, "Original", clause.original_text)
    if clause.explanation:
        _write_subsection(doc, "Why it matters", clause.explanation)
    if clause.market_comparison:
        _write_subsection(doc, "Market comparison", clause.market_comparison)
    if clause.suggested_redline:
        _write_subsection(doc, "Suggested revision", clause.suggested_redline)

    # Visual separator between clauses (but not after the last one).
    if not is_last:
        _add_thin_rule(doc)


def _write_subsection(doc: DocumentType, label: str, text: str | None) -> None:
    if not text:
        return
    doc.add_heading(label, level=3)
    for paragraph_text in _normalize_text(text).split("\n\n"):
        if paragraph_text.strip():
            doc.add_paragraph(paragraph_text.strip())


def _add_thin_rule(doc: DocumentType) -> None:
    """Insert a hairline rule between clauses. Implemented as a
    paragraph with a bottom border so it flows with the document
    rather than as a drawing object.

    OOXML requires a strict child-element order inside <w:pPr>: pBdr
    must precede spacing. We set spacing via python-docx's high-level
    API (which inserts <w:spacing>), then insert <w:pBdr> at the
    correct position rather than appending it at the end.
    """
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(12)
    rule.paragraph_format.space_after = Pt(6)

    p_pr = rule._p.get_or_add_pPr()

    p_bdr = p_pr.makeelement(qn("w:pBdr"), {})
    bottom = p_bdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "1",
            qn("w:color"): "CCCCCC",
        },
    )
    p_bdr.append(bottom)

    # pBdr must come before spacing per OOXML schema. Find the
    # spacing element (added by paragraph_format above) and insert
    # pBdr immediately before it. If for some reason spacing isn't
    # present, fall back to appending at the start of pPr.
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is not None:
        spacing.addprevious(p_bdr)
    else:
        p_pr.insert(0, p_bdr)


def _write_benign(doc: DocumentType, benign: list[Clause]) -> None:
    if not benign:
        return

    doc.add_heading("Reviewed without issue", level=1)
    intro = doc.add_paragraph("The following clauses were reviewed and no concerns were flagged.")
    intro.paragraph_format.space_after = Pt(8)

    for clause in benign:
        type_label = _format_clause_type(clause.clause_type)
        position_label = str(clause.position).zfill(3)
        doc.add_paragraph(
            f"Clause {position_label}   {type_label}",
            style="List Bullet",
        )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _dedupe_clauses(clauses: Iterable[Clause]) -> list[Clause]:
    seen: set[str] = set()
    unique: list[Clause] = []
    for clause in clauses:
        key = (clause.original_text or "").strip()
        if key in seen:
            continue
        seen.add(key)
        unique.append(clause)
    return unique


def _normalize_text(text: str) -> str:
    """Undo PDF column line breaks. Single newlines within a paragraph
    become spaces; paragraph breaks (double newlines) survive."""
    stripped = text.strip()
    normalised = re.sub(r"(?<!\n)\n(?!\n)", " ", stripped)
    normalised = re.sub(r"[ \t]+", " ", normalised)
    return normalised.strip()


def _format_clause_type(snake: str) -> str:
    if not snake:
        return "Unknown"
    return " ".join(part.capitalize() for part in snake.split("_") if part)


def _format_date(value: datetime | None) -> str:
    if value is None:
        return "date unknown"
    # Cross-platform day-without-leading-zero: format the month and year
    # normally, then strip the leading zero from the day manually. Avoids
    # the %-d (POSIX) vs %#d (Windows) split.
    day = value.day  # int, no leading zero
    return value.strftime("%B ") + f"{day}, " + value.strftime("%Y")
